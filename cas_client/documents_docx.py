"""DOCX counterparts of documents.py's loan documents (Liquidación de
Préstamo, Pagaré, Contrato, Cronograma, Comprobante de Pago) -- same data
sources (GetLoanById/GetClientById/GetAmortizationSchedule) and the same legal
clause text, but saved as an editable .docx instead of a flattened PDF so a
user can tweak specific fields (e.g. correct a client address) without
regenerating from the app.

Mirrors documents.py's structure section-by-section; reuses its shared
constants (_COMPANY_*, _TERM_*, _ESTADOS_LABEL) rather than duplicating them,
so the authorised commercial terms can't drift between the PDF and the DOCX
of the same contract."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from cas_client import assets, documents, theme
from cas_client.formatting import (
    fecha,
    fecha_hora,
    gs,
    rate_percent,
    rate_percent_mensual,
)


def _rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    )


_PRIMARY = _rgb(theme.PRIMARY)
_TEXT_PRIMARY = _rgb(theme.TEXT_PRIMARY)
_TEXT_MUTED = _rgb(theme.TEXT_MUTED)


def _add_divider(document: Document) -> None:
    """Thin horizontal rule -- python-docx has no direct API for this, so it's
    built from a paragraph's bottom border (the standard OOXML workaround)."""
    paragraph = document.add_paragraph()
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), theme.BORDER.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = _PRIMARY


def _dias_de_gracia() -> str:
    """ "11 (once) días" -- el plazo desde el que se devenga la mora, en
    dígitos y letras, igual que documents.py."""
    dias = documents._TERM_MORATORY_GRACE_DAYS
    return f"{dias} ({documents._numero_en_letras(dias)}) días"


def _cuotas_para_acelerar() -> str:
    """ "4 (cuatro) cuotas vencidas" -- el umbral que habilita a exigir el
    total adeudado."""
    cuotas = documents._TERM_ACCELERATION_INSTALLMENTS
    return f"{cuotas} ({documents._numero_en_letras(cuotas)}) cuotas vencidas"


def _add_header(document: Document, title: str) -> None:
    """Logo + legal-entity block + centered title, same layout as
    documents.py's _header() / HeaderDocumentos.png.

    Ya no recibe draft_banner: el cartel "BORRADOR -- TEXTO LEGAL PENDIENTE
    DE REVISIÓN" se quitó de los tres documentos que lo llevaban al
    autorizarse el texto legal y cargarse las condiciones reales (ver los
    _TERM_* de documents.py)."""
    table = document.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.8)

    logo_run = table.cell(0, 0).paragraphs[0].add_run()
    logo_run.add_picture(assets.LOGO_FULL_PNG, width=Inches(1.3))

    text_cell = table.cell(0, 1)
    name_run = text_cell.paragraphs[0].add_run(
        f"{documents._COMPANY_NAME} — RUC: {documents._COMPANY_RUC}"
    )
    name_run.bold = True
    name_run.font.size = Pt(12)
    name_run.font.color.rgb = _TEXT_PRIMARY

    addr_paragraph = text_cell.add_paragraph()
    addr_run = addr_paragraph.add_run(documents._COMPANY_ADDRESS)
    addr_run.font.size = Pt(9)
    addr_run.font.color.rgb = _TEXT_MUTED

    phone_paragraph = text_cell.add_paragraph()
    phone_run = phone_paragraph.add_run(f"Cel: {documents._COMPANY_PHONE}")
    phone_run.font.size = Pt(9)
    phone_run.font.color.rgb = _TEXT_MUTED

    _add_divider(document)
    _add_title(document, title)


def _add_labeled_lines(document: Document, lines: list[tuple[str, str]]) -> None:
    for label, value in lines:
        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)


def _add_client_block(document: Document, client) -> None:
    _add_labeled_lines(
        document,
        [
            ("Cliente", f"{client.first_name} {client.last_name}"),
            ("Documento", client.national_id),
            ("Dirección", client.address),
            ("Teléfono", client.phone_number),
        ],
    )


def _add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = _PRIMARY


def _add_footer_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _TEXT_MUTED


def _add_signature_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(36)
    paragraph.add_run(text)


def _add_cargos_y_garantia(document: Document, loan) -> None:
    """BR-LOAN-005/006 desglose -- same fields/condition as documents.py's
    _cargos_y_garantia_block(), purely informational (see CLAUDE.md)."""
    filas = [
        (nombre, monto)
        for nombre, monto in (
            ("Impuesto al interés", loan.charge_interest_tax),
            ("Gastos administrativos", loan.charge_admin_fee),
            ("Seguro de cancelación de deuda", loan.charge_cancellation_insurance),
            ("Seguros contratados", loan.charge_contracted_insurance),
        )
        if monto
    ]

    if filas:
        _add_section_heading(document, "Cargos y seguros")
        table = document.add_table(rows=1 + len(filas) + 1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Concepto"
        table.rows[0].cells[1].text = "Monto (Gs)"
        for row_index, (nombre, monto) in enumerate(filas, start=1):
            cells = table.rows[row_index].cells
            cells[0].text = nombre
            cells[1].text = gs(monto)
        total_cells = table.rows[-1].cells
        total_cells[0].text = "Total cargos"
        total_cells[1].text = gs(loan.total_charges)

    if loan.guarantee_type:
        _add_labeled_lines(
            document,
            [
                (
                    "Garantía",
                    f"{loan.guarantee_type} — Monto aplicado: {gs(loan.guarantee_amount)}",
                ),
                (
                    "Total del crédito (incl. cargos)",
                    gs(loan.total_credit_with_charges),
                ),
            ],
        )


def comprobante_pago_docx(loan, client, payment) -> Document:
    """DOCX de documents.py's comprobante_pago_html() -- BR-LOAN-011. Sin
    banner de borrador, mismo criterio que el Cronograma: no tiene texto legal
    a revisar, solo el detalle de un pago ya registrado."""
    document = Document()
    _add_header(document, "Comprobante de Pago")
    _add_client_block(document, client)
    _add_labeled_lines(document, [("Préstamo", loan.id)])

    filas = [
        ("Monto abonado", gs(payment.amount_paid)),
        (
            "Cuota(s) abonada(s)",
            documents.cuotas_cubiertas_texto(
                payment.covered_installments, payment.total_installments
            ),
        ),
        # Misma hora local que la versión HTML del comprobante -- ver la nota
        # en documents.comprobante_pago_html sobre por qué no es un strftime
        # directo sobre el naive UTC que devuelve ToDatetime().
        ("Fecha y hora del pago", fecha_hora(payment.paid_at.ToDatetime())),
        *documents.filas_medio_de_pago(payment),  # BR-CAJA-004
        ("Total pagado del préstamo", gs(payment.total_paid)),
        ("Saldo restante", gs(payment.remaining_balance)),
    ]
    table = document.add_table(rows=1 + len(filas), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Concepto"
    table.rows[0].cells[1].text = "Detalle"
    for row_index, (concepto, detalle) in enumerate(filas, start=1):
        cells = table.rows[row_index].cells
        cells[0].text = concepto
        cells[1].text = detalle

    if payment.status == "PAID":
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Con este pago el préstamo queda totalmente cancelado.")
        run.bold = True
        run.font.color.rgb = _rgb(theme.SUCCESS)

    _add_labeled_lines(
        document,
        [
            (
                "Registrado por",
                documents.responsable(
                    payment.recorded_by_name, payment.recorded_by_national_id
                ),
            )
        ],
    )
    _add_footer_note(
        document,
        "Comprobante emitido por el sistema de CREDIMED UME. El saldo restante "
        "puede variar por cargos o ajustes posteriores a la fecha de emisión de "
        "este comprobante.",
    )
    return document


def reporte_periodo_docx(report, generated_by: str = "") -> Document:
    """BR-DASH-002, contraparte .docx de documents.reporte_periodo_html().
    Sin banner de borrador (no tiene texto legal, solo cifras calculadas),
    igual criterio que cronograma_docx.

    report: dashboard_service_pb2.GetPeriodReportResponse"""
    document = Document()
    _add_header(document, "Reporte de Cierre de Período")

    lineas = [
        ("Período", f"del {fecha(report.start_date)} al {fecha(report.end_date)}")
    ]
    if generated_by:
        lineas.append(("Generado por", generated_by))
    _add_labeled_lines(document, lineas)

    # Misma definición de contenido que el PDF -- ver documents._filas_reporte.
    filas = documents._filas_reporte(report)
    table = document.add_table(rows=1 + len(filas), cols=3)
    table.style = "Table Grid"
    for col, text in enumerate(["Sección", "Concepto", "Valor"]):
        table.rows[0].cells[col].text = text
    for row_index, (seccion, concepto, valor) in enumerate(filas, start=1):
        cells = table.rows[row_index].cells
        cells[0].text = seccion
        cells[1].text = concepto
        cells[2].text = valor

    _add_footer_note(
        document,
        '"Movimiento" y "Cobranza" miden lo ocurrido dentro del período '
        'seleccionado. "Situación al cierre" es una foto del estado actual de '
        "la cartera al momento de generar este reporte, no del último día del "
        "período.",
    )
    _add_footer_note(
        document, "Documento generado por el sistema de CREDIMED UME. Uso interno."
    )
    return document


def liquidacion_docx(loan, client, schedule) -> Document:
    """loan: loan_service_pb2.GetLoanByIdResponse
    client: client_service_pb2.GetClientByIdResponse
    schedule: loan_service_pb2.GetAmortizationScheduleResponse"""
    document = Document()
    _add_header(document, "Liquidación de Préstamo")
    _add_client_block(document, client)

    estado = documents._ESTADOS_LABEL.get(loan.status, loan.status)
    _add_labeled_lines(
        document,
        [
            ("Préstamo", loan.id),
            ("Estado", estado),
            ("Capital", gs(loan.principal_amount)),
            (
                "Tasa de interés",
                f"{rate_percent(loan.interest_rate)} anual "
                f"({rate_percent_mensual(loan.interest_rate)} mensual sobre "
                "saldos deudores)",
            ),
            ("Plazo", f"{loan.term_months} meses"),
            ("Total pagado", gs(loan.total_paid)),
            ("Saldo restante", gs(loan.remaining_balance)),
        ],
    )

    _add_cargos_y_garantia(document, loan)

    _add_section_heading(document, "Cronograma de amortización")
    table = document.add_table(rows=1 + len(schedule.installments), cols=5)
    table.style = "Table Grid"
    for col, text in enumerate(
        ["Cuota", "Monto (Gs)", "Capital (Gs)", "Interés (Gs)", "Saldo (Gs)"]
    ):
        table.rows[0].cells[col].text = text
    for row_index, installment in enumerate(schedule.installments, start=1):
        cells = table.rows[row_index].cells
        cells[0].text = str(installment.installment_number)
        cells[1].text = gs(installment.payment_amount)
        cells[2].text = gs(installment.principal_portion)
        cells[3].text = gs(installment.interest_portion)
        cells[4].text = gs(installment.remaining_balance)

    _add_footer_note(
        document,
        "Documento generado por el sistema de CREDIMED UME. Válido únicamente junto "
        "con la firma y sello de la entidad.",
    )
    return document


def _add_garantia_line(document: Document, loan) -> None:
    """Línea de codeudor/garantía, mismo criterio que documents.py's
    _garantia_linea() (BR-LOAN-005 no distingue codeudor de garantía en
    general)."""
    if loan.guarantee_type:
        valor = f"{loan.guarantee_type} — Monto: {gs(loan.guarantee_amount)}"
    else:
        valor = "Sin garantía registrada"
    _add_labeled_lines(document, [("Garantía / Codeudor solidario", valor)])


def pagare_docx(loan, client) -> Document:
    document = Document()
    _add_header(document, "Pagaré a la Orden")
    _add_client_block(document, client)
    _add_garantia_line(document, loan)

    body = document.add_paragraph()
    body.add_run(f"DECLARO(AMOS) ADEUDAR a {documents._COMPANY_NAME} la suma de ")
    amount_run = body.add_run(f"Guaraníes {gs(loan.principal_amount)}")
    amount_run.bold = True
    body.add_run(
        ", que PAGARÉ(MOS) solidariamente, a su orden, libre de gastos y sin "
        f"protesto, en {loan.term_months} cuotas iguales, mensuales y "
        f"consecutivas, con vencimiento la primera de ellas el día "
        f"{fecha(loan.first_due_date)}, y las siguientes cuotas en esas mismas "
        f"fechas de los meses subsiguientes hasta su total cancelación, en "
        f"el domicilio de {documents._COMPANY_NAME}, sito en "
        f"{documents._COMPANY_ADDRESS}."
    )

    interest_paragraph = document.add_paragraph()
    interest_paragraph.add_run(
        "Queda expresamente pactado que los importes de las cuotas "
        "documentadas en este instrumento devengarán un interés "
        f"compensatorio del {rate_percent_mensual(loan.interest_rate)} "
        "mensual sobre saldos deudores."
    )

    mora_paragraph = document.add_paragraph()
    mora_paragraph.add_run(
        "En caso de mora se aplicará, sobre cada cuota vencida e impaga, un "
        "interés moratorio en carácter punitorio del "
        f"{documents._TERM_MORATORY_RATE}, que se devengará a partir de los "
        f"{_dias_de_gracia()} corridos contados desde la fecha de su primer "
        "vencimiento."
    )

    acceleration_paragraph = document.add_paragraph()
    acceleration_paragraph.add_run(
        f"La falta de pago de {_cuotas_para_acelerar()} facultará a "
        f"{documents._COMPANY_NAME} a exigir el total adeudado, inclusive "
        "las cuotas no vencidas, produciéndose la mora por el mero "
        "vencimiento del plazo, sin necesidad de ningún requerimiento "
        "judicial y/o extrajudicial."
    )

    jurisdiction_paragraph = document.add_paragraph()
    jurisdiction_paragraph.add_run(
        "Todas las partes intervinientes en este documento se someten a la "
        "jurisdicción y competencia de los Jueces y Tribunales de "
        f"{documents._TERM_JURISDICTION_CITY}."
    )

    _add_labeled_lines(document, [("Préstamo", loan.id)])
    _add_signature_line(document, "Firma del deudor: ______________________________")
    _add_footer_note(document, "Lugar y fecha: ______________________________")
    return document


def cronograma_docx(loan, client, schedule) -> Document:
    """DOCX de documents.py's cronograma_html() -- ver ese docstring para el
    razonamiento (documento standalone para entregar al cliente, con nombre
    del cliente y asesor responsable)."""
    document = Document()
    _add_header(document, "Cronograma de Pago")
    _add_client_block(document, client)

    asesor = documents.responsable(
        loan.created_by_full_name,
        loan.created_by_national_id,
        respaldo=loan.created_by_username,
    )
    _add_labeled_lines(
        document,
        [
            ("Asesor responsable", asesor),
            ("Préstamo", loan.id),
            ("Capital", gs(loan.principal_amount)),
            ("Tasa de interés", rate_percent(loan.interest_rate)),
            ("Plazo", f"{loan.term_months} meses"),
            ("Primer vencimiento", fecha(loan.first_due_date)),
        ],
    )

    table = document.add_table(rows=1 + len(schedule.installments), cols=6)
    table.style = "Table Grid"
    for col, text in enumerate(
        [
            "Cuota",
            "Vencimiento",
            "Monto (Gs)",
            "Capital (Gs)",
            "Interés (Gs)",
            "Saldo (Gs)",
        ]
    ):
        table.rows[0].cells[col].text = text
    for row_index, installment in enumerate(schedule.installments, start=1):
        cells = table.rows[row_index].cells
        cells[0].text = str(installment.installment_number)
        cells[1].text = fecha(installment.due_date)
        cells[2].text = gs(installment.payment_amount)
        cells[3].text = gs(installment.principal_portion)
        cells[4].text = gs(installment.interest_portion)
        cells[5].text = gs(installment.remaining_balance)

    _add_footer_note(
        document,
        "Este cronograma es informativo y está sujeto a los términos y "
        "condiciones establecidos en el Pagaré y el Contrato de Préstamo "
        "firmados. Copia entregada al cliente.",
    )
    return document


def contrato_docx(loan, client) -> Document:
    document = Document()
    _add_header(document, "Contrato de Préstamo")

    intro = document.add_paragraph()
    intro.add_run(
        f"Entre {documents._COMPANY_NAME}, con RUC {documents._COMPANY_RUC}, "
        f"con domicilio en {documents._COMPANY_ADDRESS}, en adelante "
        f"{documents._COMPANY_NAME} o LA ENTIDAD, por una parte; y por la "
        "otra el(la/los) Sr(a)(es). abajo identificado(s), en adelante "
        "EL(LOS) PRESTATARIO(S), convienen en celebrar el presente Contrato "
        "de Préstamo de Dinero, sujeto a las cláusulas y condiciones "
        "siguientes."
    )

    _add_client_block(document, client)
    _add_garantia_line(document, loan)

    _add_section_heading(document, "Cláusulas")

    clauses = [
        (
            "Primera (Objeto)",
            f"{documents._COMPANY_NAME} otorga al(los) Prestatario(s) un "
            f"préstamo de dinero por la suma de Guaraníes "
            f"{gs(loan.principal_amount)}, que se desembolsa a la firma del "
            "presente instrumento, conjuntamente con un Pagaré a la orden "
            "por dicho monto, destinado a servir como título de crédito.",
        ),
        (
            "Segunda (Reembolso)",
            "El(Los) Prestatario(s) se compromete(n) a reembolsar el "
            f"préstamo otorgado en {loan.term_months} cuotas iguales, "
            "mensuales y consecutivas, bajo el sistema de amortización "
            f"francés (cuota fija), venciendo la primera de ellas el día "
            f"{fecha(loan.first_due_date)}, mediante transferencia bancaria, "
            "débito directo o descuento en cuenta, según lo acordado.",
        ),
        (
            "Tercera (Intereses)",
            "Se acuerda el pago de un interés compensatorio del "
            f"{rate_percent_mensual(loan.interest_rate)} mensual sobre "
            "saldos deudores, abonado junto con las cuotas de amortización "
            "del capital. Para el caso de falta de pago en la fecha "
            "convenida, se aplicará además, sobre cada cuota vencida e "
            "impaga, un interés moratorio en carácter punitorio del "
            f"{documents._TERM_MORATORY_RATE}, que se devengará a partir de "
            f"los {_dias_de_gracia()} corridos contados desde la fecha de su "
            "primer vencimiento.",
        ),
        (
            "Cuarta (Mora y vencimiento anticipado)",
            "La mora se producirá por el mero vencimiento de los plazos, "
            "sin necesidad de interpelación judicial alguna. La falta de "
            f"pago de {_cuotas_para_acelerar()} hará decaer de pleno "
            "derecho todos los plazos estipulados para el pago, facultando "
            f"a {documents._COMPANY_NAME} a declarar vencidas todas las "
            "cuotas y exigir el pago de la totalidad de la deuda, "
            "ejecutando para el efecto el Pagaré suscripto junto con este "
            "contrato.",
        ),
        (
            "Quinta (Central de riesgo crediticio)",
            "El(Los) Prestatario(s) autoriza(n) expresamente a "
            f"{documents._COMPANY_NAME} para que, por su cuenta o a través "
            "de terceros, recabe información sobre su situación "
            "patrimonial y crediticia, y para que, en caso de atraso en el "
            "pago, informe sus datos a las centrales de riesgo crediticio "
            "correspondientes, conforme a la legislación vigente.",
        ),
        (
            "Sexta (Jurisdicción)",
            "Todas las partes intervinientes en este contrato se someten a "
            "la jurisdicción y competencia de los Jueces y Tribunales de "
            f"{documents._TERM_JURISDICTION_CITY}.",
        ),
    ]
    for heading, text in clauses:
        paragraph = document.add_paragraph()
        heading_run = paragraph.add_run(f"{heading}: ")
        heading_run.bold = True
        paragraph.add_run(text)

    _add_labeled_lines(document, [("Préstamo", loan.id)])
    _add_signature_line(document, "Firma del deudor: ______________________________")
    _add_footer_note(
        document, f"Firma de {documents._COMPANY_NAME}: ______________________________"
    )
    return document
